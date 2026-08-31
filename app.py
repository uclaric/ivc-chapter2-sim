from flask import Flask, render_template, request, jsonify, session, send_file
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.units import inch
import os
import json
import random
import re
import unicodedata
import secrets
from io import BytesIO
from xml.sax.saxutils import escape
from werkzeug.middleware.proxy_fix import ProxyFix
import logging

from config.chapter2 import (
    CHAPTER, CORE_CAST, GUESTS, QUESTION_POOL, TEXTBOOK_KNOWLEDGE, BASE_INSTRUCTIONS
)
from config.guest_audio import GUEST_AUDIO

app = Flask(__name__)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)

IS_PRODUCTION = os.environ.get("SIM_PRODUCTION", "0") == "1"
SIM_SECRET_KEY = os.environ.get("SIM_SECRET_KEY", "dev-only-change-me")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
SIM_INSTRUCTOR_TOKEN = os.environ.get("SIM_INSTRUCTOR_TOKEN", "")

app.secret_key = SIM_SECRET_KEY
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=IS_PRODUCTION,
    MAX_CONTENT_LENGTH=12 * 1024 * 1024,
)

logging.basicConfig(level=os.environ.get("LOG_LEVEL", "INFO"))
logger = logging.getLogger("ivc-sim")

if IS_PRODUCTION and SIM_SECRET_KEY == "dev-only-change-me":
    logger.warning("SIM_SECRET_KEY is using the development fallback. Set a strong secret before student use.")
if not OPENAI_API_KEY:
    logger.warning("OPENAI_API_KEY is not set. OpenAI-powered routes will fail until it is configured.")

client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
MODEL = os.environ.get("OPENAI_MODEL", "gpt-5.5")
TTS_MODEL = os.environ.get("OPENAI_TTS_MODEL", "gpt-4o-mini-tts")
TRANSCRIBE_MODEL = os.environ.get("OPENAI_TRANSCRIBE_MODEL", "gpt-4o-mini-transcribe")

CORE_AUDIO = {
    "Prof. Epps": {
        "voice": "onyx",
        "instructions": (
            "Warm, confident American male college professor in his late 60s with a slightly deep voice. "
            "Relaxed live coffee-table conversation, expressive pauses, varied pacing, occasional dry humor, never narration."
        ),
    },
    "Sophia": {"voice": "marin", "instructions": "Young American woman in college. Thoughtful, friendly, warm, contemporary and naturally expressive."},
    "Ethan": {"voice": "ash", "instructions": "Young American male college student. Friendly, curious, intelligent, easygoing, lightly self-deprecating."},
    "Carlos": {"voice": "cedar", "instructions": "Young Hispanic male from Mexicali. Fluent English with a subtle northern Mexican/Baja California accent. Relaxed, warm, contemporary. Never British or caricatured."},
    "Aaliyah": {"voice": "coral", "instructions": "Young Black American woman in college. Warm, confident, thoughtful, socially aware, direct when appropriate, natural and contemporary."},
    "Freja": {"voice": "marin", "instructions": "Young Swedish woman speaking fluent English with a light natural Swedish accent. Friendly, curious, socially conscious. Never exaggerated."},
}
CHARACTER_AUDIO = {**CORE_AUDIO, **GUEST_AUDIO}
ALLOWED_VOICES = {"alloy", "ash", "ballad", "coral", "echo", "fable", "nova", "onyx", "sage", "shimmer", "verse", "marin", "cedar"}

VOICE_PREVIEW_TEXT = {
    "Prof. Epps": "All right, everybody. Congress. Five hundred thirty-five ambitious people and somehow the furniture survives.",
    "Sophia": "Congress looks simple in a diagram until you ask who actually controls the agenda.",
    "Ethan": "So if Congress was designed to be slow, how do we know when slow turns into broken?",
    "Carlos": "I want to know what all these rules mean for regular people waiting for something to actually get done.",
    "Aaliyah": "Representation sounds good, but I want to know whose voices actually make it into the room.",
    "Freja": "The Senate is fascinating to me because equal states and equal citizens are not quite the same thing.",
    "James Madison": "A representative government must contend with faction without surrendering liberty to its cure.",
    "Henry Clay": "A legislature that cannot bargain will soon discover that principle alone does not keep a union together.",
    "Lyndon B. Johnson": "You can give the finest speech in Washington, but if you haven't counted the votes, you haven't done the job.",
    "Shirley Chisholm": "Representation means more than being allowed into the room. It means having the independence to speak once you are there.",
    "John McCain": "Party matters, but there are moments when the institution and your own judgment have to matter more.",
}


GUEST_PARTICIPATION_RULES = """
SPECIAL GUEST PARTICIPATION — LOCKED RULES FOR ALL GUESTS

1. Guests are active discussants, not decorative cameos.
2. After the core social welcome and introductions are complete, but BEFORE the first graded Chapter 2 question begins, Prof. Epps formally welcomes the Chapter 2 guests who will participate in that session.
3. At that moment, the student-facing room should transition to the guest-inclusive Chapter 2 group image.
4. One or two recurring classmates may react naturally to the guests' period clothing or presence using current student language when it fits the moment, for example "Okay, that outfit is fire" or "Cool threads." Keep this brief, natural, and never disrespectful.
5. Each participating guest gives only a brief introduction or opening remark. Do not turn the guest welcome into a lecture or roll call.
6. Guests may respond to Prof. Epps, classmates, one another, and the real IVC student.
7. When a guest participates meaningfully in an exchange, the guest should USUALLY turn the conversation back to the real IVC student with a direct, natural question such as asking what the student thinks, how the student would decide, what tradeoff the student sees, or how the guest's historical experience compares with the student's view.
8. A guest's question must sound like that person and arise from the current topic and what has actually been said. Do not mechanically repeat "What do you think?" every time.
9. Guests may challenge the student respectfully, ask for clarification, invite comparison, or press for a reason or example. Keep the difficulty appropriate for an introductory American Government course.
10. The real IVC student remains central. Guest-to-guest and guest-to-classmate exchanges are encouraged, but do not let extended guest dialogue crowd the real student out.
11. Do not have every guest speak on every question. Use the guest or guests whose documented perspective genuinely improves that exchange.
12. Historical and contemporary guest dialogue must be grounded ONLY in documented material attributable to that person: writings, speeches, interviews, letters, testimony, official records, public statements, or authorized autobiographical material.
13. The Sim may paraphrase documented ideas for clarity and accessibility, but MUST NOT invent a guest's beliefs, motives, private thoughts, experiences, quotations, or positions that are not supported by the historical record.
14. Never fabricate quotations. If exact wording is uncertain, paraphrase rather than quote.
15. Guests should speak in a manner appropriate to their era and documented public voice without becoming caricatures or confusing students with unnecessarily archaic language.
16. At the end of the Sim, Prof. Epps thanks the guests for participating and for sharing perspectives rooted in what they actually lived, wrote, said, or authorized.
17. The closing should briefly remind students that the Sim dramatizes documented historical perspectives and is not inventing biography.
"""

SIM_INSTRUCTIONS = BASE_INSTRUCTIONS + "\n\n" + TEXTBOOK_KNOWLEDGE + "\n\n" + GUEST_PARTICIPATION_RULES


def require_openai():
    if client is None:
        return jsonify({"error": "The Sim server is not connected to OpenAI yet."}), 503
    return None


def build_question_plan():
    accessible = random.choice(QUESTION_POOL["accessible"])
    medium = random.sample(QUESTION_POOL["medium"], 2)
    challenging = random.choice(QUESTION_POOL["challenging"])
    plan = [
        {**accessible, "difficulty": "accessible"},
        {**medium[0], "difficulty": "medium"},
        {**medium[1], "difficulty": "medium"},
        {**challenging, "difficulty": "challenging"},
    ]
    # Keep the intended difficulty ramp, while the actual wording and speaker remain dynamic.
    return plan


def current_question_state():
    plan = session.get("question_plan") or []
    idx = int(session.get("question_index", 0))
    current = plan[idx] if idx < len(plan) else None
    nxt = plan[idx + 1] if idx + 1 < len(plan) else None
    return plan, idx, current, nxt


def guest_context_for(question):
    if not question:
        return "No new historical guest is required."
    names = question.get("guest_tags", [])[:3]
    chunks = []
    for name in names:
        spec = GUESTS.get(name)
        if not spec:
            continue
        chunks.append(
            f"{name}: era={spec['era']}; expertise={', '.join(spec['expertise'])}; speaking rule={spec['speech_style']}"
        )
    return "\n".join(chunks) or "No new historical guest is required."


def parse_sim_payload(response):
    raw = (response.output_text or "").strip()
    allowed = set(CORE_CAST) | set(GUESTS.keys()) | {"Professor Epps"}
    try:
        payload = json.loads(raw)
        if not isinstance(payload, dict):
            raise ValueError("Payload must be an object.")
        turns = payload.get("turns", [])
        if not isinstance(turns, list):
            turns = []
        cleaned_turns = []
        for turn in turns:
            if not isinstance(turn, dict):
                continue
            speaker = str(turn.get("speaker", "Prof. Epps")).strip()
            text = str(turn.get("text", "")).strip()
            expression = str(turn.get("expression", "")).strip().lower()
            if not text:
                continue

            # Normalize the model's natural full-name variants to the short
            # display names used by the UI/expression map. Without this, names
            # such as "Sophia Martinez" fail validation and fall back to
            # Prof. Epps, making every portrait/label appear as Epps.
            speaker_aliases = {
                "Professor Epps": "Prof. Epps",
                "Ric Epps": "Prof. Epps",
                "Professor Ric Epps": "Prof. Epps",
                "Sophia Martinez": "Sophia",
                "Ethan Williams": "Ethan",
                "Carlos Rodriguez": "Carlos",
                "Aaliyah Brooks": "Aaliyah",
                "Freja Lindström": "Freja",
                "Freja Lindstrom": "Freja",
                "Lyndon Johnson": "Lyndon B. Johnson",
                "LBJ": "Lyndon B. Johnson",
                "Madison": "James Madison",
                "Clay": "Henry Clay",
                "Chisholm": "Shirley Chisholm",
                "McCain": "John McCain",
            }
            speaker = speaker_aliases.get(speaker, speaker)
            if speaker not in allowed:
                speaker = "Prof. Epps"
            if expression not in {"neutral", "smile", "amused", "skeptical", "thinking", "surprised"}:
                expression = ""
            cleaned_turns.append({"speaker": speaker, "text": text[:4000], "action": "", "expression": expression})
        if not cleaned_turns:
            raise ValueError("No valid turns returned.")
        qresult = str(payload.get("question_result", "in_progress")).strip().lower()
        if qresult not in {"not_started", "in_progress", "answered", "clarify", "skipped"}:
            qresult = "in_progress"
        return {"turns": cleaned_turns, "wait_for_student": bool(payload.get("wait_for_student", True)), "question_result": qresult}
    except Exception:
        return {
            "turns": [{"speaker": "Prof. Epps", "text": raw or "Give me one second. The room lost its train of thought.", "action": "", "expression": "thinking"}],
            "wait_for_student": True,
            "question_result": "in_progress",
        }


@app.get("/healthz")
def healthz():
    return jsonify({
        "status": "ok",
        "service": CHAPTER["service"],
        "chapter": CHAPTER["number"],
        "openai_configured": bool(OPENAI_API_KEY),
        "model": MODEL,
        "tts_model": TTS_MODEL,
        "transcribe_model": TRANSCRIBE_MODEL,
    })


@app.route("/")
def index():
    return render_template("index.html")


@app.post("/api/start")
def api_start():
    data = request.get_json(force=True)
    name = (data.get("name") or "").strip()
    origin = (data.get("origin") or "").strip()
    fact = (data.get("fact") or "").strip()
    if not name:
        return jsonify({"error": "Name is required."}), 400
    missing = require_openai()
    if missing:
        return missing

    name, origin, fact = name[:80], origin[:160], fact[:800]
    session.clear()
    session["student_name"] = name
    session["question_plan"] = build_question_plan()
    session["question_index"] = 0
    session["question_started"] = False
    session["question_results"] = []
    session["sim_complete"] = False

    user_input = f'''
The real student has just entered the Chapter 2 discussion room and introduced themself:
Name: {name}
From: {origin or "not specified"}
Crazy fun fact: {fact or "not specified"}

Complete the social opening and guest welcome, but do NOT begin the graded Congress discussion yet.
- Prof. Epps reacts briefly and freshly to the student's fun fact.
- Sophia, Ethan, Carlos, Aaliyah, and Freja each introduce themselves naturally.
- Prof. Epps introduces himself last.
- Natural interjections are welcome; do not make it feel like roll call.
- After those introductions, Prof. Epps formally welcomes the Chapter 2 special guests who will participate in this session.
- Choose the guests from the Chapter 2 guest roster whose perspectives are most useful for the planned discussion. All five do not have to speak immediately.
- One or two recurring students may make a brief, respectful modern reaction to the guests' period clothing or presence.
- Each welcomed guest gives only a short introductory remark rooted in documented historical material.
- Do not ask the first graded question yet.
- End at a natural point where {name} can respond to the guest arrival/welcome if they wish.
Set question_result to "not_started".
'''
    try:
        response = client.responses.create(model=MODEL, instructions=SIM_INSTRUCTIONS, input=user_input)
        session["previous_response_id"] = response.id
        return jsonify(parse_sim_payload(response))
    except Exception:
        logger.exception("OpenAI start request failed")
        return jsonify({"error": "The discussion room could not start just now. Please try again."}), 502


@app.post("/api/reply")
def api_reply():
    data = request.get_json(force=True)
    student_text = (data.get("text") or "").strip()
    if not student_text:
        return jsonify({"error": "Response is required."}), 400
    missing = require_openai()
    if missing:
        return missing

    student_text = student_text[:2500]
    name = session.get("student_name", "Student")
    previous_response_id = session.get("previous_response_id")
    if not previous_response_id:
        return jsonify({"error": "This discussion session has expired. Please restart the Sim from the beginning."}), 409

    plan, idx, current, nxt = current_question_state()
    if not current:
        # Defensive closure if the client somehow calls again after completion.
        return jsonify({
            "turns": [{"speaker": "Prof. Epps", "text": f"You're all set, {name}. The Chapter 2 Sim-Discussion is complete. Head back to Canvas for the remaining assignment instructions.", "action": "", "expression": "amused"}],
            "wait_for_student": False,
            "question_result": "answered",
            "complete": True,
            "sim_complete": True,
        })

    started = bool(session.get("question_started", False))
    current_guests = guest_context_for(current)
    next_summary = nxt["target"] if nxt else "There is no next question; close the Sim after this target is completed or skipped."

    if not started:
        task_block = f'''
The social opening and formal guest welcome are complete. {name}'s latest line is social conversation, NOT an answer to a graded question.
Now transition naturally into the FIRST graded Congress discussion target.
CURRENT HIDDEN TARGET: {current['target']}
Appropriate historical guest options for this target:
{current_guests}
Use the already welcomed guest whose documented perspective best fits this target. If another guest becomes more relevant later, rotate naturally.
The guest may comment briefly, question another participant, or directly ask {name} what they think. A guest who contributes substantively should usually bring {name} into the exchange with a direct, natural question tied to the topic.
Ask the target in fresh conversational wording appropriate to an accessible opening. Do not reveal difficulty or question number.
Set question_result to "in_progress".
'''
    else:
        task_block = f'''
CURRENT HIDDEN TARGET ID: {current['id']}
CURRENT HIDDEN TARGET: {current['target']}
CURRENT DIFFICULTY (never reveal): {current['difficulty']}
Appropriate historical guest options:
{current_guests}
NEXT TARGET SUMMARY (use only if current target becomes answered/skipped): {next_summary}

Evaluate {name}'s latest response ONLY against the current target.
- If it is a meaningful attempt that sufficiently engages the target, set question_result="answered". Briefly respond/synthesize, then if another target remains, transition naturally toward it and ask it without revealing question numbers or difficulty. You may rotate guests if the next topic warrants it.
- If the answer is thin but salvageable, set question_result="in_progress" and give one gentle, inviting push for a reason, example, distinction, consequence, or one more step of thought.
- If {name} asks for clarification or says the question is not understood, set question_result="clarify". Rephrase the SAME target more simply and give only neutral context, not the answer.
- If {name} explicitly says skip/pass/does not want to answer, set question_result="skipped". Accept immediately and without judgment. If another target remains, move naturally to it. Do not ask them to reconsider.
- If this is the fourth target and it becomes answered or skipped, close with a concise synthesis. Prof. Epps must thank the participating guests for sharing perspectives grounded in what they lived, wrote, said publicly, or authorized in autobiographical accounts; briefly remind the student that the guest dialogue is based on documented historical material; then clearly say the Chapter 2 Sim-Discussion is complete.
'''

    user_input = f'''
The student's name is {name}.
{name} just said:
{student_text}

{task_block}

Keep the room conversational. Prof. Epps does not automatically respond first. Historical guests, recurring classmates, and Prof. Epps may speak to each other.
When a historical guest contributes meaningfully, that guest should usually ask {name} a direct, natural question or respond to something {name} previously said and invite them back into the exchange.
Do not force every guest into every topic and do not let guest-to-guest dialogue crowd out the real student.
Usually give 2-4 short turns, then wait for {name}, except the final completion may close without another question.
'''

    try:
        response = client.responses.create(model=MODEL, instructions=SIM_INSTRUCTIONS, previous_response_id=previous_response_id, input=user_input)
        session["previous_response_id"] = response.id
        payload = parse_sim_payload(response)

        if not started:
            session["question_started"] = True
        else:
            result = payload.get("question_result")
            if result in {"answered", "skipped"}:
                results = session.get("question_results", [])
                results.append({"id": current["id"], "result": result, "difficulty": current["difficulty"]})
                session["question_results"] = results
                session["question_index"] = idx + 1
                session["question_started"] = True  # response may already have introduced the next target
                if idx + 1 >= len(plan):
                    session["sim_complete"] = True
                    payload["complete"] = True
                    payload["sim_complete"] = True
                    payload["wait_for_student"] = False
                    visible_text = " ".join(str(t.get("text", "")) for t in payload.get("turns", [])).lower()
                    if "sim-discussion is complete" not in visible_text and "sim discussion is complete" not in visible_text:
                        payload["turns"].append({
                            "speaker": "Prof. Epps",
                            "text": f"All right, {name}, that completes your Chapter 2 Sim-Discussion. Before you submit, you can choose whether you'd like a brief preliminary assessment of how you did.",
                            "action": "",
                            "expression": "amused",
                        })
            elif result == "clarify":
                session["question_started"] = True

        payload["progress"] = min(int(session.get("question_index", 0)), 4)
        payload["sim_complete"] = bool(session.get("sim_complete", False))
        # Hidden grading metadata is returned for the client to persist, but the UI does not display it.
        payload["grading_marker"] = {
            "question_id": current["id"],
            "difficulty": current["difficulty"],
            "result": payload.get("question_result"),
            "zero_if_skipped": payload.get("question_result") == "skipped",
        }
        return jsonify(payload)
    except Exception:
        logger.exception("OpenAI reply request failed")
        return jsonify({"error": "The room had trouble responding. Your typed response is still visible; please try again."}), 502


ASSESSMENT_INSTRUCTIONS = """
You are generating a preliminary learning assessment for a completed POLS C1000 Chapter 2 Sim-Discussion.
This is NOT the professor's final grade and must never include points, percentages, letter grades, rubric scores, or numerical scores.

Choose exactly one level:
- Excellent / Great Work
- Good / Solid Work
- Fair / Okay
- Struggled a Bit — Consider Redoing
- You Should Seriously Consider Redoing

Base the assessment only on the student's participation visible in the transcript, considering:
1. meaningful participation and engagement,
2. accurate understanding of Chapter 2 Congress and representation course material,
3. reasoning, evidence, tradeoffs, contradictions, and response to other perspectives,
4. clear communication in the student's own words.

Pay particular attention to the Chapter 2 themes:
- representation and democratic accountability,
- districts, elections, redistricting, gerrymandering, and political incentives,
- lawmaking, leadership, committees, Senate/House procedures, compromise, and gridlock,
- money, lobbying, party pressure, polarization, institutional norms, and constitutional duty.

Return valid JSON only:
{
  "level": "one allowed level",
  "summary": "one concise sentence",
  "points": ["3 to 5 concise observations"]
}

Be specific about strengths and areas to improve. If the student struggled, identify the concepts or reasoning problems clearly and recommend redoing the Sim when appropriate. Do not be harsh, vague, or patronizing.
"""


@app.post("/api/assessment")
def api_assessment():
    if not session.get("sim_complete"):
        return jsonify({"error": "The Sim must be completed before a preliminary assessment is available."}), 409

    missing = require_openai()
    if missing:
        return missing

    data = request.get_json(force=True, silent=True) or {}
    transcript_text = str(data.get("transcript") or "").strip()
    if not transcript_text:
        return jsonify({"error": "No transcript was received for assessment."}), 400
    if len(transcript_text) > 250_000:
        return jsonify({"error": "That transcript is too large to assess."}), 413

    try:
        response = client.responses.create(
            model=MODEL,
            instructions=ASSESSMENT_INSTRUCTIONS,
            input=transcript_text,
        )
        raw = (response.output_text or "").strip()
        payload = json.loads(raw)

        allowed_levels = {
            "Excellent / Great Work",
            "Good / Solid Work",
            "Fair / Okay",
            "Struggled a Bit — Consider Redoing",
            "You Should Seriously Consider Redoing",
        }
        level = str(payload.get("level") or "").strip()
        if level not in allowed_levels:
            raise ValueError("Invalid assessment level.")

        summary = str(payload.get("summary") or "").strip()[:800]
        points = payload.get("points") or []
        if not isinstance(points, list):
            points = []
        cleaned_points = [str(p).strip()[:600] for p in points if str(p).strip()][:5]
        if len(cleaned_points) < 3:
            raise ValueError("Assessment needs at least three observations.")

        return jsonify({
            "level": level,
            "summary": summary,
            "points": cleaned_points,
            "disclaimer": "This is a preliminary learning assessment, not your course grade. Prof. Epps independently grades your complete Sim-Discussion using the course rubric."
        })
    except Exception:
        logger.exception("Preliminary assessment generation failed")
        return jsonify({"error": "The preliminary assessment could not be generated just now. You may continue to submission without it."}), 502


def instructor_token_valid(candidate):
    if not SIM_INSTRUCTOR_TOKEN or not candidate:
        return False
    try:
        return secrets.compare_digest(str(candidate), SIM_INSTRUCTOR_TOKEN)
    except Exception:
        return False


@app.get("/api/instructor/status")
def api_instructor_status():
    token = request.args.get("token", "")
    return jsonify({"authorized": instructor_token_valid(token)})


@app.post("/api/instructor/skip")
def api_instructor_skip():
    data = request.get_json(force=True, silent=True) or {}
    token = str(data.get("token") or "")
    if not instructor_token_valid(token):
        return jsonify({"error": "Instructor test mode is not authorized."}), 403

    plan = session.get("question_plan") or build_question_plan()
    session["question_plan"] = plan
    session["question_index"] = len(plan)
    session["question_started"] = True
    session["question_results"] = [
        {"id": q.get("id", f"q{i+1}"), "result": "answered", "difficulty": q.get("difficulty", "")}
        for i, q in enumerate(plan)
    ]
    session["sim_complete"] = True
    name = session.get("student_name", "Professor")

    return jsonify({
        "turns": [{
            "speaker": "Prof. Epps",
            "text": f"Instructor test mode: jumping to the Chapter 2 completion workflow for {name}.",
            "action": "",
            "expression": "amused",
        }],
        "wait_for_student": False,
        "question_result": "answered",
        "progress": 4,
        "complete": True,
        "sim_complete": True,
    })


@app.post("/api/transcribe")
def api_transcribe():
    missing = require_openai()
    if missing:
        return missing
    audio = request.files.get("audio")
    if audio is None:
        return jsonify({"error": "No microphone recording was received."}), 400
    raw = audio.read()
    if not raw:
        return jsonify({"error": "The microphone recording was empty."}), 400
    if len(raw) > 10 * 1024 * 1024:
        return jsonify({"error": "That recording is too large. Please try a shorter response."}), 413
    filename = (audio.filename or "student-response.webm")[:120]
    mimetype = audio.mimetype or "audio/webm"
    try:
        transcript = client.audio.transcriptions.create(
            model=TRANSCRIBE_MODEL,
            file=(filename, raw, mimetype),
            language="en",
            prompt=(
                "American Government college discussion about Congress and representation. Preserve the student's wording. "
                "Expect terms such as Congress, House, Senate, representation, delegate, trustee, politico, redistricting, "
                "gerrymandering, committee, filibuster, oversight, appropriation, lobbying, polarization, and Constitution."
            ),
        )
        text = (getattr(transcript, "text", "") or "").strip()
        if not text:
            return jsonify({"error": "I couldn't make out enough speech to transcribe. Please try again."}), 422
        return jsonify({"text": text[:2500]})
    except Exception:
        logger.exception("Student speech transcription failed")
        return jsonify({"error": "The microphone recording could not be transcribed just now. You can still type your response."}), 502


@app.post("/api/speech")
def api_speech():
    missing = require_openai()
    if missing:
        return missing
    data = request.get_json(force=True)
    speaker = (data.get("speaker") or "Prof. Epps").strip()
    text = (data.get("text") or "").strip()
    requested_voice = (data.get("voice") or "").strip()
    preview = bool(data.get("preview"))
    if preview and not text:
        text = VOICE_PREVIEW_TEXT.get(speaker, "Welcome to the Sim-Discussion.")
    if not text:
        return jsonify({"error": "Text is required."}), 400
    text = text[:4000]
    profile = CHARACTER_AUDIO.get(speaker, {"voice": "alloy", "instructions": "Speak naturally and conversationally with clear expressive inflection."})
    voice = requested_voice if requested_voice in ALLOWED_VOICES else profile["voice"]
    try:
        speech = client.audio.speech.create(model=TTS_MODEL, voice=voice, input=text, instructions=profile["instructions"], response_format="mp3")
        return send_file(BytesIO(speech.read()), mimetype="audio/mpeg", as_attachment=False, download_name="sim-dialogue.mp3")
    except Exception:
        logger.exception("Speech generation failed")
        return jsonify({"error": "Speech generation failed. Text dialogue is still available."}), 502


def safe_student_filename(name):
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", (name or "Student").strip())
    cleaned = cleaned.strip("._-")
    return cleaned[:80] or "Student"


def export_payload():
    data = request.get_json(force=True, silent=True) or {}
    transcript_text = str(data.get("transcript") or "").strip()
    student_name = str(data.get("student_name") or session.get("student_name") or "Student").strip()

    if not transcript_text:
        return None, None, (jsonify({"error": "No transcript was received for export."}), 400)
    if len(transcript_text) > 250_000:
        return None, None, (jsonify({"error": "That transcript is too large to export."}), 413)

    return transcript_text, student_name[:120], None


def normalize_pdf_text(value):
    replacements = {
        "’": "'", "‘": "'", "“": '"', "”": '"',
        "—": "-", "–": "-", "…": "...", "•": "-",
        "→": "->", "←": "<-", "·": "-", "🎙️": "", "📋": "",
        "📄": "", "🔊": "", "↻": "", "Ⅱ": "II",
    }
    value = str(value or "")
    for source, target in replacements.items():
        value = value.replace(source, target)
    return unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")


@app.post("/api/export/docx")
def api_export_docx():
    transcript_text, student_name, error = export_payload()
    if error:
        return error

    try:
        output = BytesIO()
        document = Document()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(11)

        lines = transcript_text.splitlines()
        for index, raw_line in enumerate(lines):
            line = raw_line.rstrip()

            if not line:
                document.add_paragraph("")
                continue

            if index == 0:
                p = document.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(15)
                continue

            if line in {"CHAPTER 2 SIM-DISCUSSION", "Congress and Representation: Who Speaks for Whom?"}:
                p = document.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
                continue

            if set(line) == {"-"}:
                document.add_paragraph("────────────────────────────────────────")
                continue

            if line.endswith(":") or " — STUDENT RESPONSE:" in line:
                p = document.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                continue

            document.add_paragraph(line)

        document.save(output)
        output.seek(0)

        filename = f"Chapter_2_Sim_Discussion_{safe_student_filename(student_name)}.docx"
        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )
    except Exception:
        logger.exception("DOCX transcript export failed")
        return jsonify({"error": "The DOCX file could not be created just now. You can still copy the transcript."}), 500


@app.post("/api/export/pdf")
def api_export_pdf():
    transcript_text, student_name, error = export_payload()
    if error:
        return error

    try:
        output = BytesIO()
        doc = SimpleDocTemplate(
            output,
            pagesize=letter,
            rightMargin=0.6 * inch,
            leftMargin=0.6 * inch,
            topMargin=0.6 * inch,
            bottomMargin=0.6 * inch,
            title="Chapter 2 Sim-Discussion Transcript",
            author="POLS C1000",
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "SimTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=8,
        )
        subhead_style = ParagraphStyle(
            "SimSubhead",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            spaceBefore=4,
            spaceAfter=4,
        )
        body_style = ParagraphStyle(
            "SimBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            spaceAfter=5,
        )
        speaker_style = ParagraphStyle(
            "SimSpeaker",
            parent=body_style,
            fontName="Helvetica-Bold",
            spaceBefore=6,
            spaceAfter=2,
        )

        story = []
        lines = transcript_text.splitlines()
        for index, raw_line in enumerate(lines):
            line = normalize_pdf_text(raw_line).strip()

            if not line:
                story.append(Spacer(1, 5))
                continue

            if index == 0:
                story.append(Paragraph(escape(line), title_style))
                continue

            if line in {"CHAPTER 2 SIM-DISCUSSION", "Congress and Representation: Who Speaks for Whom?"}:
                story.append(Paragraph(escape(line), subhead_style))
                continue

            if set(line) == {"-"}:
                story.append(Spacer(1, 4))
                story.append(HRFlowable(width="100%", thickness=0.7, color="#777777"))
                story.append(Spacer(1, 4))
                continue

            escaped_line = escape(line)
            if line.endswith(":") or " - STUDENT RESPONSE:" in line:
                story.append(Paragraph(escaped_line, speaker_style))
            else:
                story.append(Paragraph(escaped_line, body_style))

        doc.build(story)
        output.seek(0)

        filename = f"Chapter_2_Sim_Discussion_{safe_student_filename(student_name)}.pdf"
        return send_file(
            output,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=filename,
        )
    except Exception:
        logger.exception("PDF transcript export failed")
        return jsonify({"error": "The PDF file could not be created just now. You can still copy the transcript."}), 500


@app.post("/api/reset")
def api_reset():
    session.clear()
    return jsonify({"ok": True})


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=not IS_PRODUCTION)
